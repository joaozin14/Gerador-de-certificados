import pandas as pd
from tkinter import *
from tkinter import messagebox
from tkinter import ttk
from docx import Document
from docx.shared import Pt

def filtrar():
    for linha in range(len(dados_usuarios)):

        todas_linhas = treeview_dados.get_children()
        treeview_dados.delete(*todas_linhas)
        if cpf_entry.get() == "":
            botao_pesquisa.config(text="Filtrar")
            for linha in range(len(dados_usuarios)):
                data_inicio_ano = dados_usuarios.iloc[linha, 3].split("-")[0] 
                data_inicio_mes = dados_usuarios.iloc[linha, 3].split("-")[1] 
                data_inicio_dia = dados_usuarios.iloc[linha, 3].split("-")[2] 
                data_inicio_certa = f"{data_inicio_dia}/{data_inicio_mes}/{data_inicio_ano}"

                data_fim_ano = dados_usuarios.iloc[linha, 4].split("-")[0] 
                data_fim_mes = dados_usuarios.iloc[linha, 4].split("-")[1] 
                data_fim_dia = dados_usuarios.iloc[linha, 4].split("-")[2] 
                data_fim_certa = f"{data_fim_dia}/{data_fim_mes}/{data_fim_ano}"

                treeview_dados.insert("", "end",
                                values=(str(dados_usuarios.iloc[linha, 0]), #CPF
                                        str(dados_usuarios.iloc[linha, 1]), #NOME
                                        str(dados_usuarios.iloc[linha, 2]), #RG
                                        str(data_inicio_certa), #DATA INICIO
                                        str(data_fim_certa), #DATA FIM
                                        str(dados_usuarios.iloc[linha, 5]), #EMAIL
                                    ))

        else:
            botao_pesquisa.config(text="Limpar filtro")
            for linha in range(len(dados_usuarios)):

                data_inicio_ano = dados_usuarios.iloc[linha, 3].split("-")[0] 
                data_inicio_mes = dados_usuarios.iloc[linha, 3].split("-")[1] 
                data_inicio_dia = dados_usuarios.iloc[linha, 3].split("-")[2] 
                data_inicio_certa = f"{data_inicio_dia}/{data_inicio_mes}/{data_inicio_ano}"

                data_fim_ano = dados_usuarios.iloc[linha, 4].split("-")[0] 
                data_fim_mes = dados_usuarios.iloc[linha, 4].split("-")[1] 
                data_fim_dia = dados_usuarios.iloc[linha, 4].split("-")[2] 
                data_fim_certa = f"{data_fim_dia}/{data_fim_mes}/{data_fim_ano}"

                if cpf_entry.get() == str(dados_usuarios.iloc[linha, 0]):
                    treeview_dados.insert("", "end",
                            values=(str(dados_usuarios.iloc[linha, 0]), #CPF
                                    str(dados_usuarios.iloc[linha, 1]), #NOME
                                    str(dados_usuarios.iloc[linha, 2]), #RG
                                    str(data_inicio_certa), #DATA INICIO
                                    str(data_fim_certa), #DATA FIM
                                    str(dados_usuarios.iloc[linha, 5]), #EMAIL
                                    ))


def gerar_certificado():
    arquivo_word = Document("Certificado.docx")
    estilo = arquivo_word.styles["Normal"]

    cpf_aluno = cpf_entry.get()
    nome_aluno = nome_entry.get()
    data_inicio = dtinicio_entry.get()
    data_fim = dtfim_entry.get()
    nome_instrutor = "Clevison Santos"

    frase_parte1 = f"Portador do CPF {cpf_aluno} Concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de"
    frasae_montada = f"{frase_parte1} {data_inicio} á {data_fim}."

    for paragrafo in arquivo_word.paragraphs:
        if "@nome" in paragrafo.text:
            paragrafo.text = nome_aluno
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

        if "@DataFim" in paragrafo.text:
            paragrafo.text = frasae_montada
            fonte = estilo.font
            fonte.name = "Calibri (Corpo)"
            fonte.size = Pt(24)

    caminho_certificado = f"C:\\Users\\joaoh\\OneDrive\\Área de Trabalho\\Certificado_python\\{nome_aluno}.docx"
    arquivo_word.save(caminho_certificado)

    cpf_entry.delete(0, END)
    nome_entry.delete(0, END)
    rg_entry.delete(0, END)
    dtinicio_entry.delete(0, END)
    dtfim_entry.delete(0, END)
    email_entry.delete(0, END)

    messagebox.showinfo("Mensagem", "Certificado gerado com sucesso!")


def gerar_massa():
    for linha_informacoes in treeview_dados.get_children():
        coluna = treeview_dados.item(linha_informacoes)["values"]

        cpf_separado = coluna[0]
        nome_separado = coluna[1]
        rg_separado = coluna[2]
        datainicio_separado = coluna[3]
        datafim_separado = coluna[4]
        instrutor_separado = "Clevison Santos"

        arquivo_word = Document("Certificado.docx")
        estilo = arquivo_word.styles["Normal"]

        frase_parte1 = f"Portador do CPF {cpf_separado} Concluiu com sucesso o curso de Python RPA, com a carga horária de 20 horas, promovido pela escola de Cursos Online de"
        frasae_montada = f"{frase_parte1} {datainicio_separado} á {datafim_separado}."

        for paragrafo in arquivo_word.paragraphs:
            if "@nome" in paragrafo.text:
                paragrafo.text = nome_separado
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

            if "@DataFim" in paragrafo.text:
                paragrafo.text = frasae_montada
                fonte = estilo.font
                fonte.name = "Calibri (Corpo)"
                fonte.size = Pt(24)

        caminho_certificado = f"C:\\Users\\joaoh\\OneDrive\\Área de Trabalho\\Certificado_python\\{nome_separado}.docx"
        arquivo_word.save(caminho_certificado)

    messagebox.showinfo("Mensagem", "Certificados gerados com sucesso!")

def double_click(Event):
    item = treeview_dados.selection()
    for i in item:
        cpf_entry.delete(0, END)
        nome_entry.delete(0, END)
        rg_entry.delete(0, END)
        dtinicio_entry.delete(0, END)
        dtfim_entry.delete(0, END)
        email_entry.delete(0, END)

        cpf_entry.insert(0, treeview_dados.item(i, "values") [0])
        nome_entry.insert(0, treeview_dados.item(i, "values") [1])
        rg_entry.insert(0, treeview_dados.item(i, "values") [2])
        dtinicio_entry.insert(0, treeview_dados.item(i, "values") [3])
        dtfim_entry.insert(0, treeview_dados.item(i, "values") [4])
        email_entry.insert(0, treeview_dados.item(i, "values") [5])


janela = Tk()
janela.title("Gerador de certificados")
stilo = ttk.Style()
stilo.theme_use("alt")
stilo.configure(".", font=("Arial 15 bold"), rowheight=30)

treeview_dados = ttk.Treeview(janela, columns=(1, 2, 3, 4, 5, 6), show="headings")
treeview_dados.column("1", anchor=CENTER)
treeview_dados.heading("1", text="CPF")

treeview_dados.column("2", anchor=CENTER)
treeview_dados.heading("2", text="NOME")

treeview_dados.column("3", anchor=CENTER)
treeview_dados.heading("3", text="RG")

treeview_dados.column("4", anchor=CENTER)
treeview_dados.heading("4", text="DATA INÍCIO")

treeview_dados.column("5", anchor=CENTER)
treeview_dados.heading("5", text="DATA FIM")

treeview_dados.column("6", anchor=CENTER)
treeview_dados.heading("6", text="EMAIL")

treeview_dados.grid(row=4, column=0, columnspan=6, sticky="NSEW", pady=15)

treeview_dados.bind("<Double-1>", double_click)


dados_usuarios = pd.read_excel("Dados.xlsx")
dados_usuarios["Data Inicio"] = dados_usuarios["Data Inicio"].astype(str) #Convertendo a coluna "DATA INICIO" para texto
dados_usuarios["Data Fim"] = dados_usuarios["Data Fim"].astype(str)


#Lendo linha por linha no excel, feito isso, insere na treeview, transforma em string, daí com o iloc, pega o que está na variavel "linha" através de seu indice 
for linha in range(len(dados_usuarios)):
    data_inicio_ano = dados_usuarios.iloc[linha, 3].split("-")[0] 
    data_inicio_mes = dados_usuarios.iloc[linha, 3].split("-")[1] 
    data_inicio_dia = dados_usuarios.iloc[linha, 3].split("-")[2] 
    data_inicio_certa = f"{data_inicio_dia}/{data_inicio_mes}/{data_inicio_ano}"

    data_fim_ano = dados_usuarios.iloc[linha, 4].split("-")[0] 
    data_fim_mes = dados_usuarios.iloc[linha, 4].split("-")[1] 
    data_fim_dia = dados_usuarios.iloc[linha, 4].split("-")[2] 
    data_fim_certa = f"{data_fim_dia}/{data_fim_mes}/{data_fim_ano}"

    treeview_dados.insert("", "end",
                          values=(str(dados_usuarios.iloc[linha, 0]), #CPF
                                  str(dados_usuarios.iloc[linha, 1]), #NOME
                                  str(dados_usuarios.iloc[linha, 2]), #RG
                                  str(data_inicio_certa), #DATA INICIO
                                  str(data_fim_certa), #DATA FIM
                                  str(dados_usuarios.iloc[linha, 5]), #EMAIL
                                  ))
    
cpf_label = Label(text="CPF", font="Arial 12")
cpf_label.grid(row=0, column=0, sticky="E", pady=15)
cpf_entry = Entry(font="Arial 12")
cpf_entry.grid(row=0, column=1, sticky="W", pady=15)

nome_label = Label(text="Nome", font="Arial 12")
nome_label.grid(row=0, column=2, sticky="E", pady=15)
nome_entry = Entry(font="Arial 12")
nome_entry.grid(row=0, column=3, sticky="W", pady=15)

rg_label = Label(text="RG", font="Arial 12")
rg_label.grid(row=0, column=4, sticky="E", pady=15)
rg_entry = Entry(font="Arial 12")
rg_entry.grid(row=0, column=5, sticky="W", pady=15)

dtinicio_label = Label(text="Data Inicío", font="Arial 12")
dtinicio_label.grid(row=1, column=0, sticky="E", pady=15)
dtinicio_entry = Entry(font="Arial 12")
dtinicio_entry.grid(row=1, column=1, sticky="W", pady=15)

dtfim_label = Label(text="Data Fim", font="Arial 12")
dtfim_label.grid(row=1, column=2, sticky="E", pady=15)
dtfim_entry = Entry(font="Arial 12")
dtfim_entry.grid(row=1, column=3, sticky="W", pady=15)

email_label = Label(text="Email", font="Arial 12")
email_label.grid(row=1, column=4, sticky="E", pady=15)
email_entry = Entry(font="Arial 12")
email_entry.grid(row=1, column=5, sticky="W", pady=15)

botao_pesquisa = Button(text="PESQUISAR", font="Arial 14", command=filtrar)
botao_pesquisa.grid(row=5, column=0, columnspan=2, sticky="NSEW", padx=20)

botao_certificado = Button(text="GERAR CERTIFICADO", font="Arial 14", command=gerar_certificado)
botao_certificado.grid(row=5, column=2, columnspan=2, sticky="NSEW", padx=20)

botao_massa = Button(text="GERAR CERTIFICADO EM MASSA", font="Arial 14", command=gerar_massa)
botao_massa.grid(row=5, column=4, columnspan=2, sticky="NSEW", padx=20)

janela.mainloop()